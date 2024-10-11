import os
import fitz  # install via PyMuPDF
from docx import Document
import docx
from logger import logger

class DocumentProcessor:
    def __init__(self, folder_path):
        self.folder_path = folder_path

    def get_document_files(self):
        try:
            supported_extensions = ('.pdf', '.docx', '.txt')
            document_files = [
                os.path.join(self.folder_path, file)
                for file in os.listdir(self.folder_path)
                if file.lower().endswith(supported_extensions)
            ]
            logger.info(f"Found {len(document_files)} document files.")
            return document_files
        except Exception as e:
            logger.error(f"Error accessing folder {self.folder_path}: {e}")
            return []

    def extract_text(self, doc_path):
        try:
            _, ext = os.path.splitext(doc_path)
            ext = ext.lower()
            text = ""

            if ext == '.pdf':
                text = self.extract_text_pdf(doc_path)
            elif ext == '.docx':
                text = self.extract_text_docx(doc_path)
            elif ext == '.txt':
                text = self.extract_text_txt(doc_path)
            else:
                logger.warning(f"Unsupported file extension for {doc_path}")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {doc_path}: {e}")
            return ""

    def extract_text_pdf(self, pdf_path):
        try:
            logger.info(f"Opening PDF: {pdf_path}")
            doc = fitz.open(pdf_path)
            text = []
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text()
                if page_text:
                    text.append(page_text)
                if page_num % 10 == 0:
                    logger.info(f"Extracted text from page {page_num}")
            doc.close()
            full_text = "\n".join(text)
            logger.info(f"Extracted text from {pdf_path} with total length {len(full_text)} characters.")
            return full_text
        except Exception as e:
            logger.error(f"Error extracting text from {pdf_path}: {e}")
            return ""

    def extract_text_docx(self, docx_path):
        try:
            logger.info(f"Opening DOCX: {docx_path}")
            doc = docx.Document(docx_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            logger.info(f"Extracted text from {docx_path} with total length {len(text)} characters.")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {docx_path}: {e}")
            return ""

    def extract_text_txt(self, txt_path):
        try:
            logger.info(f"Opening TXT: {txt_path}")
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            logger.info(f"Extracted text from {txt_path} with total length {len(text)} characters.")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {txt_path}: {e}")
            return ""

    def get_num_pages(self, doc_path):
        try:
            _, ext = os.path.splitext(doc_path)
            ext = ext.lower()

            if ext == '.pdf':
                return self.get_num_pages_pdf(doc_path)
            elif ext == '.docx':
                return self.get_num_pages_docx(doc_path)
            elif ext == '.txt':
                return self.get_num_pages_txt(doc_path)
            else:
                logger.warning(f"Unsupported file extension for {doc_path}")
                return 1
        except Exception as e:
            logger.error(f"Error getting number of pages for {doc_path}: {e}")
            return 1  # Default to 1 page if unable to determine

    def get_num_pages_pdf(self, pdf_path):
        try:
            logger.info(f"Getting number of pages for PDF: {pdf_path}")
            doc = fitz.open(pdf_path)
            num_pages = doc.page_count
            doc.close()
            logger.info(f"{pdf_path} has {num_pages} pages.")
            return num_pages
        except Exception as e:
            logger.error(f"Error getting number of pages for {pdf_path}: {e}")
            return 1  # Default to 1 page if unable to determine

    def get_num_pages_docx(self, docx_path):
        try:
            logger.info(f"Getting number of pages for DOCX: {docx_path}")
            # DOCX doesn't have pages, so we can estimate based on word count
            doc = docx.Document(docx_path)
            word_count = sum(len(para.text.split()) for para in doc.paragraphs)
            # Assuming 300 words per page
            num_pages = max(1, word_count // 300)
            logger.info(f"{docx_path} has estimated {num_pages} pages based on word count.")
            return num_pages
        except Exception as e:
            logger.error(f"Error getting number of pages for {docx_path}: {e}")
            return 1

    def get_num_pages_txt(self, txt_path):
        try:
            logger.info(f"Getting number of pages for TXT: {txt_path}")
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            word_count = len(text.split())
            num_pages = max(1, word_count // 300)
            logger.info(f"{txt_path} has estimated {num_pages} pages based on word count.")
            return num_pages
        except Exception as e:
            logger.error(f"Error getting number of pages for {txt_path}: {e}")
            return 1
